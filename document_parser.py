"""
Document Parser Module

This module handles parsing different document formats (text, PDF)
and extracting their content for date processing.
"""

import os
import logging
from io import BytesIO

logger = logging.getLogger(__name__)

def parse_document(file_path):
    """
    Parse a document and extract its text content based on file type.
    
    Args:
        file_path (str): Path to the document
        
    Returns:
        str: Extracted text content
        
    Raises:
        ValueError: If file type is unsupported
        FileNotFoundError: If file does not exist
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.txt':
        return parse_text_file(file_path)
    elif file_extension in ['.pdf', '.PDF']:
        return parse_pdf_file(file_path)
    elif file_extension in ['.docx', '.doc']:
        return parse_word_file(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def parse_text_file(file_path):
    """
    Parse a text file and extract its content.
    
    Args:
        file_path (str): Path to the text file
        
    Returns:
        str: Text content
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with a different encoding if UTF-8 fails
        with open(file_path, 'r', encoding='latin-1') as file:
            return file.read()


def parse_pdf_file(file_path):
    """
    Parse a PDF file and extract its text content.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text content
    """
    try:
        import PyPDF2
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            
            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n\n"
            
            return text
    except ImportError:
        logger.error("PyPDF2 library not found. Install with: pip install PyPDF2")
        raise
    except Exception as e:
        logger.error(f"Error parsing PDF file: {e}")
        raise


def parse_word_file(file_path):
    """
    Parse a Microsoft Word document and extract its text content.
    
    Args:
        file_path (str): Path to the Word document
        
    Returns:
        str: Extracted text content
    """
    try:
        import docx
        
        doc = docx.Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text
    except ImportError:
        logger.error("python-docx library not found. Install with: pip install python-docx")
        raise
    except Exception as e:
        logger.error(f"Error parsing Word document: {e}")
        raise


def extract_text_from_file(file_path):
    """
    High-level function to extract text from a file regardless of type.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        str: Extracted text content
    """
    logger.info(f"Extracting text from file: {file_path}")
    
    try:
        text = parse_document(file_path)
        logger.info(f"Successfully extracted {len(text)} characters from {file_path}")
        return text
    except Exception as e:
        logger.error(f"Failed to extract text from {file_path}: {e}")
        raise